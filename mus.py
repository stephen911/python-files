import os
import datetime
# h = int(input("hour"))
# m = int(input("minutes"))
# ampm = input("am or pm")
# if ampm == "pm":
#     h = h + 12
# while True:
#     if (h == datetime.datetime.now().hour) and (m == datetime.datetime.now().minute):
#         print("wake up")
#         break
#
# path = r"E:\music\format\\"
# files = os.listdir(path)
# #os.startfile(r"C:\Users\Stephen Dapaah\PycharmProjects\hello\screenshot1.png")
import pyautogui as p
# #p.hotkey("fn", "f11")
p.press("playpause")
# #p.press("pause")
# #p.press("prevtrack")
# #p.press("nexttrack")
# #p.press("volumeup")
# #p.press("volumeup")
# #p.press("volumeup")
# import pyowm
# owm = pyowm.OWM("73afb0146dc0eaaf6b2d92789a4a0ffd")
# # search for current whather in london(great britain)
# observation = owm.weather_at_place("Accra")
# w = observation.get_weather()
#
# #Wheather details
# city = "accra"
# print(w.get_wind())
# print(w.get_humidity())
# print(w.get_rain)
# print(w.get_snow)
# sr = w.get_sunrise_time(timeformat="iso")
# ss = w.get_sunset_time(timeformat="iso")
# loc = owm.three_hours_forecast(city)
# clouds = str(loc.will_have_rain())
# print(sr, ss, loc, clouds)
# temp = w.get_temperature(unit="celsius")
# for key, val in temp.items():
#     print(f"{key} => {val}")
# num = "A God Like You"
# for file in files:
#     #print(file)
#     s = file.replace("_", " ")
#     r = s.replace("-", " ")
#     if num in r:
#         #print(r)
#         now = r.replace(" ", "-")
#         p = path + "{}".format(now)
#         #print(p)
#         #os.startfile(p)
#     #print(file.replace(" ", "-"))
#     #n = file.replace(" ", "-")
#     #now = n.split(".")
#     #print(now)
#     #current = now[-2]
#     #print(current)
#
#     # if "com" in current:
#     #     #print(current)
#     #     c = file.replace(current, "")
#     #     v = c.replace("..", ".")
#     #     print(v)
#     #os.replace(path + "\\{}".format(file), path + "\\{}".format(n))
#
#     #final = current[-2]
#     # os.replace(path)
#


    #n = file[-4:]
    #m = file.replace(n, ".mp3")
    #n = file[-4:]
    #m = file.replace(n, ".png")
    #os.replace(r"C:\Users\Stephen Dapaah\Desktop\snip\{}".format(file), r"C:\Users\Stephen Dapaah\Desktop\snip\{}".format(m))
    #os.replace(path + "\\{}".format(file), path + "\\{}".format(n))
# from docx import *
# document = opendocx(r'test.docx')
# words = document.xpath('//w:r', namespaces=document.nsmap)
#
# WPML_URI = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
# tag_rPr = WPML_URI + 'rPr'
# tag_highlight = WPML_URI + 'highlight'
# tag_val = WPML_URI + 'val'
#
# for word in words:
#     for rPr in word.findall(tag_rPr):
#         if rPr.find(tag_highlight).attrib[tag_val] == 'yellow':
#             #print(word.find(tag_t).text)

from docx import Document
from docx.shared import Inches

document = Document()

document.add_heading('Document Title', 0)

p = document.add_paragraph('A plain paragraph having some ')
p.add_run('bold').bold = True
p.add_run(' and some ')
p.add_run('italic.').italic = True

document.add_heading('Heading, level 1', level=1)
document.add_paragraph('Intense quote', style='Intense Quote')

document.add_paragraph(
    'first item in unordered list', style='List Bullet'
)
document.add_paragraph(
    'first item in ordered list', style='List Number'
)

document.add_picture('monty-truth.png', width=Inches(1.25))

records = (
    (3, '101', 'Spam'),
    (7, '422', 'Eggs'),
    (4, '631', 'Spam, spam, eggs, and spam')
)

table = document.add_table(rows=1, cols=3)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Qty'
hdr_cells[1].text = 'Id'
hdr_cells[2].text = 'Desc'
for qty, id, desc in records:
    row_cells = table.add_row().cells
    row_cells[0].text = str(qty)
    row_cells[1].text = id
    row_cells[2].text = desc

document.add_page_break()

document.save('demo.docx')